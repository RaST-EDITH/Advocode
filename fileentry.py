import os
import docx
from docx2pdf import convert
# from docx.enum.text import WD_ALIGN_PARAGRAPH

def fillBlanks( file, cred ) :
    doc = docx.Document( file )

    paragraphs = doc.paragraphs
    for para in paragraphs :
        newpara = list( para.text.split(" ") )
        for i in range(len(newpara)) :
            if newpara[i] in cred.keys() :
                newpara[i] = cred[newpara[i]]

        para.clear()
        para.add_run( " ".join(newpara))
    path = file.split(".docx")[0]
    path = path + f'-{cred["idName"]}.docx'
    doc.save( path )
    convert( path )
    if os.path.exists( path ):
        os.remove( path )
    return
