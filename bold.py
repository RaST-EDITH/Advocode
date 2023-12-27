import os
import docx
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH

def changeBold( target_word ) :

    doc = docx.Document(r'Licence-to-use-Copyright2.docx')
    new_doc = docx.Document()
    count = 0

    # Iterate through paragraphs to find and format the word
    for paragraph in doc.paragraphs:
        if ( target_word in paragraph.text ) or count == 0:

            if count == 0 :
                new_paragraph = new_doc.add_paragraph()
                new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = new_paragraph.add_run(paragraph.text)
                run.bold = True
                count = 1
            
            else :
                # Split the paragraph text into segments based on the target word
                segments = paragraph.text.split(target_word)

                # Create a new paragraph for the modified text
                new_paragraph = new_doc.add_paragraph()
                for i, segment in enumerate(segments):
                    if segment:
                        new_paragraph.add_run(segment)  # Add the surrounding text
                    if  ( i < len(segments) - 1 ) :
                        run = new_paragraph.add_run(target_word)  # Add the target word
                        run.bold = True
                        run.underline = True
